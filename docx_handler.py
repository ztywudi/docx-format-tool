"""
底稿整理助手 v1.1.0 - docx 处理核心模块
纯离线，处理 Word 文档的样式提取与应用
"""

import os
import json
import shutil
import logging
from datetime import datetime
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# 配置日志
logger = logging.getLogger(__name__)


# ========== 字体与尺寸常量 ==========

CN_FONT_SIZES = {
    "初号": 42, "小初": 36, "一号": 26, "小一": 24,
    "二号": 22, "小二": 18, "三号": 16, "小三": 15,
    "四号": 14, "小四": 12, "五号": 10.5, "小五": 9,
    "六号": 7.5, "小六": 6.5,
}

PT_TO_CN = {v: k for k, v in CN_FONT_SIZES.items()}

CN_FONTS = ["宋体", "仿宋", "黑体", "楷体", "微软雅黑", "方正小标宋", "方正仿宋", "仿宋_GB2312"]
EN_FONTS = ["Times New Roman", "Arial", "Calibri"]

ALIGN_MAP = {
    "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
    "居中": WD_ALIGN_PARAGRAPH.CENTER,
    "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
    "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

ALIGN_REVERSE = {
    WD_ALIGN_PARAGRAPH.LEFT: "左对齐",
    WD_ALIGN_PARAGRAPH.CENTER: "居中",
    WD_ALIGN_PARAGRAPH.RIGHT: "右对齐",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐",
}

LINE_SPACING_OPTIONS = ["1.0", "1.25", "1.5", "2.0", "固定值"]

PAGE_NUM_FORMATS = ["1, 2, 3...", "-1-, -2-...", "第1页", "第1页 / 共N页"]


def _get_today_formats():
    """生成基于当前日期的格式选项，确保默认值始终有效"""
    today = datetime.now()
    return [
        f"{today.year}年{today.month}月{today.day}日",
        f"{today.year}.{today.month:02d}.{today.day:02d}",
        f"{today.year}-{today.month:02d}-{today.day:02d}",
        f"{today.year}/{today.month:02d}/{today.day:02d}",
    ]


DATE_FORMATS = _get_today_formats()

TEMPLATE_CATEGORIES = ["审计报告", "税审报告", "投标文件", "通用模板", "其他"]

# 正文首行缩进换算：1个中文字符 ≈ 12pt（小四号）≈ 2字符缩进 ≈ 0.74cm
# 精确换算：字符数 × 字号(pt) × 12700 = Emu值
CHAR_TO_EMU_FACTOR = 12700  # 1pt = 12700 EMU


# ========== 工具函数 ==========

def pt_to_cn_size(pt_value):
    """将磅值转换为最接近的中文字号"""
    if pt_value is None:
        return "五号"
    pt_value = float(pt_value)
    closest = min(CN_FONT_SIZES.values(), key=lambda x: abs(x - pt_value))
    return PT_TO_CN.get(closest, f"{pt_value:.1f}pt")


def cn_to_pt(cn_size):
    """将中文字号转换为磅值"""
    if cn_size in CN_FONT_SIZES:
        return CN_FONT_SIZES[cn_size]
    try:
        return float(cn_size.replace("pt", ""))
    except (ValueError, AttributeError):
        return 12  # 默认小四


def get_default_template():
    """返回默认模板（通用格式）"""
    today = datetime.now()
    return {
        "name": "通用模板",
        "source": "内置默认",
        "cover": {
            "auto_generate": False,
            "title_font": "宋体",
            "title_size": "一号",
            "title_bold": True,
            "title_align": "居中",
            "company_font": "宋体",
            "company_size": "三号",
            "date_font": "宋体",
            "date_size": "四号",
        },
        "headings": {
            "heading1": {
                "font": "黑体",
                "size": "三号",
                "bold": True,
                "align": "居中",
                "space_before": 12,
                "space_after": 6,
            },
            "heading2": {
                "font": "黑体",
                "size": "四号",
                "bold": True,
                "align": "左对齐",
                "space_before": 8,
                "space_after": 4,
            },
            "heading3": {
                "font": "仿宋",
                "size": "小四",
                "bold": True,
                "align": "左对齐",
                "space_before": 6,
                "space_after": 3,
            },
        },
        "body": {
            "font": "仿宋",
            "size": "小四",
            "bold": False,
            "line_spacing": 1.5,
            "line_spacing_rule": "multiple",
            "first_line_indent": 2,
            "align": "两端对齐",
        },
        "table": {
            "header_bold": True,
            "header_shading": True,
            "header_shading_color": "D9E2F3",
            "content_align": "居中",
            "auto_fit": True,
        },
        "image": {
            "center_align": True,
            "max_width_cm": 15,
        },
        "header": {
            "text": "",
            "font": "宋体",
            "size": "小五",
            "align": "居中",
            "different_first_page": True,
        },
        "footer": {
            "page_num_format": "第1页 / 共N页",
            "align": "居中",
            "start_from": 1,
            "font": "宋体",
            "size": "小五",
            "show_on_first_page": False,
        },
        "toc": {
            "include": False,
        },
        "signature": {
            "align_right": True,
            "include_date": True,
            "date_format": f"{today.year}年{today.month}月{today.day}日",
        },
    }


# ========== 样式提取（从模板 docx）==========

def _get_para_align(paragraph):
    """获取段落对齐方式"""
    if paragraph.alignment is not None:
        return ALIGN_REVERSE.get(paragraph.alignment, "左对齐")
    return "左对齐"


def _get_run_font_info(run):
    """获取 run 的字体信息"""
    font_name = None
    size_pt = None
    bold = None

    if run.font.name:
        font_name = run.font.name

    if run.font.size:
        size_pt = run.font.size.pt

    if run.font.bold is not None:
        bold = run.font.bold

    return font_name, size_pt, bold


def extract_styles_from_docx(filepath):
    """从 docx 文件中提取格式，返回模板字典"""
    doc = Document(filepath)
    template = get_default_template()
    template["source"] = f"导入自: {os.path.basename(filepath)}"

    # ---- 提取标题样式 ----
    for heading_id in [1, 2, 3]:
        style_name = f"Heading {heading_id}"
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            key = f"heading{heading_id}"

            if style.font and style.font.name:
                template["headings"][key]["font"] = style.font.name
            if style.font and style.font.size:
                template["headings"][key]["size"] = pt_to_cn_size(style.font.size.pt)
            if style.font and style.font.bold is not None:
                template["headings"][key]["bold"] = style.font.bold

            pf = style.paragraph_format
            if pf.alignment is not None:
                template["headings"][key]["align"] = ALIGN_REVERSE.get(pf.alignment, "左对齐")
            if pf.space_before:
                template["headings"][key]["space_before"] = pf.space_before.pt
            if pf.space_after:
                template["headings"][key]["space_after"] = pf.space_after.pt

    # ---- 如果样式提取不完整，从实际段落中采样 ----
    heading_samples = {1: [], 2: [], 3: []}
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("Heading"):
            for hid in [1, 2, 3]:
                if para.style.name == f"Heading {hid}":
                    heading_samples[hid].append(para)

    for hid in [1, 2, 3]:
        key = f"heading{hid}"
        if heading_samples[hid]:
            sample = heading_samples[hid][0]
            if sample.runs:
                fn, sz, bd = _get_run_font_info(sample.runs[0])
                if fn:
                    template["headings"][key]["font"] = fn
                if sz:
                    template["headings"][key]["size"] = pt_to_cn_size(sz)
                if bd is not None:
                    template["headings"][key]["bold"] = bd
            template["headings"][key]["align"] = _get_para_align(sample)

    # ---- 提取正文样式 ----
    body_keywords = ["摘要", "引言", "概述", "背景", "正文", "报告", "说明"]
    body_paras = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if len(text) > 30:
            body_paras.append(para)
        elif any(kw in text for kw in body_keywords) and len(text) > 5:
            body_paras.append(para)

    if not body_paras:
        for para in doc.paragraphs:
            if para.text.strip() and len(para.text.strip()) > 10:
                body_paras.append(para)
                if len(body_paras) >= 5:
                    break

    if body_paras:
        fonts = []
        sizes = []
        for p in body_paras:
            if p.runs:
                fn, sz, _ = _get_run_font_info(p.runs[0])
                if fn:
                    fonts.append(fn)
                if sz:
                    sizes.append(sz)

        if fonts:
            most_font = max(set(fonts), key=fonts.count)
            template["body"]["font"] = most_font

        if sizes:
            most_size = max(set(sizes), key=sizes.count)
            template["body"]["size"] = pt_to_cn_size(most_size)

        pf = body_paras[0].paragraph_format
        if pf.line_spacing:
            template["body"]["line_spacing"] = pf.line_spacing
        if pf.first_line_indent:
            # 将 cm 值换算为字符数（按小四号 12pt 估算）
            indent_cm = pf.first_line_indent.cm
            template["body"]["first_line_indent"] = round(indent_cm / 0.74)

    # ---- 提取页眉 ----
    try:
        section = doc.sections[0]
        header = section.header
        if header and header.paragraphs:
            header_text = ""
            for p in header.paragraphs:
                header_text += p.text
            if header_text.strip():
                template["header"]["text"] = header_text.strip()
                if header.paragraphs[0].runs:
                    fn, sz, _ = _get_run_font_info(header.paragraphs[0].runs[0])
                    if fn:
                        template["header"]["font"] = fn
                    if sz:
                        template["header"]["size"] = pt_to_cn_size(sz)
                template["header"]["align"] = _get_para_align(header.paragraphs[0])
    except Exception as e:
        logger.warning(f"提取页眉样式时出错：{e}")

    # ---- 提取页脚/页码 ----
    try:
        footer = section.footer
        if footer and footer.paragraphs:
            for p in footer.paragraphs:
                if "PAGE" in p.text or "页" in p.text or any(c.isdigit() for c in p.text):
                    template["footer"]["align"] = _get_para_align(p)
                    break
    except Exception as e:
        logger.warning(f"提取页脚样式时出错：{e}")

    # ---- 提取表格格式 ----
    if doc.tables:
        table = doc.tables[0]
        if table.rows:
            first_row = table.rows[0]
            for cell in first_row.cells:
                for p in cell.paragraphs:
                    if p.runs and p.runs[0].font.bold:
                        template["table"]["header_bold"] = True
                    if p.alignment:
                        template["table"]["content_align"] = ALIGN_REVERSE.get(p.alignment, "居中")

    # ---- 提取落款 ----
    for para in doc.paragraphs:
        text = para.text.strip()
        if any(kw in text for kw in ["公章", "盖章", "单位", "公司", "事务所", "日期", "年", "月"]):
            if para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                template["signature"]["align_right"] = True
            break

    return template


# ========== 样式应用（套模板）==========

def _set_font(run, font_name=None, size_pt=None, bold=None, color=None):
    """设置 run 的字体"""
    if font_name:
        run.font.name = font_name
        # 设置中文字体
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
    if size_pt and size_pt > 0:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _set_paragraph_format(paragraph, align=None, space_before=None, space_after=None,
                          line_spacing=None, line_spacing_rule=None, first_line_indent=None,
                          font_size_pt=12):
    """设置段落格式
    
    Args:
        first_line_indent: 首行缩进字符数（整数），会根据字号自动换算为实际宽度
        font_size_pt: 正文字号（磅值），用于首行缩进的精确换算
    """
    pf = paragraph.paragraph_format
    if align and align in ALIGN_MAP:
        pf.alignment = ALIGN_MAP[align]
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if first_line_indent is not None and first_line_indent > 0:
        # 按字符数精确换算：字符数 × 字号pt × 12700 = EMU
        indent_emu = int(first_line_indent * font_size_pt * CHAR_TO_EMU_FACTOR)
        pf.first_line_indent = Emu(indent_emu)


def _add_page_field(paragraph, font_name, font_size):
    """向段落中插入 PAGE 域代码（页码）"""
    # begin
    fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_begin = paragraph.add_run()
    run_begin._element.append(fld_begin)

    # instrText
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_instr = paragraph.add_run()
    run_instr._element.append(instr)

    # end
    fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_end = paragraph.add_run()
    run_end._element.append(fld_end)


def _add_numpages_field(paragraph, font_name, font_size):
    """向段落中插入 NUMPAGES 域代码（总页数）"""
    fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_begin = paragraph.add_run()
    run_begin._element.append(fld_begin)

    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
    run_instr = paragraph.add_run()
    run_instr._element.append(instr)

    fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_end = paragraph.add_run()
    run_end._element.append(fld_end)


def _apply_heading_style(doc, heading_key, settings):
    """应用标题样式"""
    style_name = f"Heading {heading_key[-1]}"
    try:
        style = doc.styles[style_name]
    except KeyError:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    sz = cn_to_pt(settings["size"])
    f = style.font
    f.name = settings["font"]
    f.size = Pt(sz)
    f.bold = settings["bold"]

    # 中文字体
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), settings["font"])

    pf = style.paragraph_format
    if settings["align"] in ALIGN_MAP:
        pf.alignment = ALIGN_MAP[settings["align"]]
    if settings.get("space_before"):
        pf.space_before = Pt(settings["space_before"])
    if settings.get("space_after"):
        pf.space_after = Pt(settings["space_after"])

    # 遍历文档段落，确保所有使用该样式的段落都被应用
    for para in doc.paragraphs:
        if para.style and para.style.name == style_name:
            _set_paragraph_format(
                para,
                align=settings["align"],
                space_before=settings.get("space_before"),
                space_after=settings.get("space_after"),
            )
            for run in para.runs:
                _set_font(run, settings["font"], sz, settings["bold"])


def _apply_body_style(doc, settings):
    """应用正文样式"""
    try:
        style = doc.styles["Normal"]
    except KeyError:
        style = doc.styles.add_style("Normal", WD_STYLE_TYPE.PARAGRAPH)

    sz = cn_to_pt(settings["size"])
    f = style.font
    f.name = settings["font"]
    f.size = Pt(sz)
    f.bold = settings.get("bold", False)

    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), settings["font"])

    pf = style.paragraph_format
    pf.alignment = ALIGN_MAP.get(settings["align"], WD_ALIGN_PARAGRAPH.JUSTIFY)
    pf.line_spacing = settings["line_spacing"]
    if settings.get("first_line_indent"):
        # 使用字符数精确换算
        indent_emu = int(settings["first_line_indent"] * sz * CHAR_TO_EMU_FACTOR)
        pf.first_line_indent = Emu(indent_emu)

    # 收集页眉/页脚段落 id，用于排除
    header_footer_para_ids = set()
    try:
        for sec in doc.sections:
            if sec.header:
                for p in sec.header.paragraphs:
                    header_footer_para_ids.add(id(p))
            if sec.footer:
                for p in sec.footer.paragraphs:
                    header_footer_para_ids.add(id(p))
    except Exception as e:
        logger.warning(f"收集页眉页脚段落时出错：{e}")

    # 遍历所有非标题段落
    heading_styles = {f"Heading {i}" for i in range(1, 4)}
    for para in doc.paragraphs:
        is_heading = para.style and para.style.name in heading_styles
        is_hf = id(para) in header_footer_para_ids

        if not is_heading and not is_hf:
            if para.runs:
                for run in para.runs:
                    _set_font(run, settings["font"], sz, settings.get("bold", False))
            # 不对正文段落强制对齐（保留原有的居中对齐等特殊情况）


def _apply_table_format(doc, settings):
    """格式化文档中的所有表格"""
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    if i == 0 and settings["header_bold"]:
                        for run in para.runs:
                            run.font.bold = True

                    if settings.get("content_align") in ALIGN_MAP:
                        para.alignment = ALIGN_MAP[settings["content_align"]]

        # 表头底纹
        if settings.get("header_shading") and table.rows:
            for cell in table.rows[0].cells:
                tcPr = cell._element.get_or_add_tcPr()
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{settings.get("header_shading_color", "D9E2F3")}" w:val="clear"/>'
                )
                tcPr.append(shading)

        # 设置表格边框
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'</w:tblBorders>'
        )
        old_borders = tblPr.find(qn('w:tblBorders'))
        if old_borders is not None:
            tblPr.remove(old_borders)
        tblPr.append(borders)


def _apply_image_format(doc, settings):
    """格式化文档中的图片：居中 + 限制最大宽度"""
    max_width_cm = settings.get("max_width_cm", 15)
    max_width_emu = int(Cm(max_width_cm))

    for para in doc.paragraphs:
        drawings = para._element.findall(qn('w:drawing'))
        if not drawings:
            continue

        # 居中对齐
        if settings.get("center_align"):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 限制图片最大宽度
        for drawing in drawings:
            # 查找 extent（图片尺寸）
            inlines = drawing.findall(qn('wp:inline'))
            anchors = drawing.findall(qn('wp:anchor'))
            for parent_elem in inlines + anchors:
                extent = parent_elem.find(qn('wp:extent'))
                if extent is not None:
                    cx = int(extent.get('cx', '0'))
                    cy = int(extent.get('cy', '0'))
                    if cx > max_width_emu and cx > 0:
                        # 等比缩放
                        ratio = max_width_emu / cx
                        new_cx = max_width_emu
                        new_cy = int(cy * ratio)
                        extent.set('cx', str(new_cx))
                        extent.set('cy', str(new_cy))
                        logger.debug(f"图片缩放：{cx}→{new_cx}, {cy}→{new_cy}")


def _apply_header_footer(doc, settings):
    """应用页眉页脚和页码"""
    if not doc.sections:
        return

    for section in doc.sections:
        # ---- 页眉 ----
        header_settings = settings.get("header", {})
        header_text = header_settings.get("text", "")
        header_align = header_settings.get("align", "居中")
        header_font = header_settings.get("font", "宋体")
        header_size = cn_to_pt(header_settings.get("size", "小五"))

        section.different_first_page_header_footer = header_settings.get("different_first_page", True)

        if header_text.strip():
            for header_type in ['default', 'first']:
                header = section.header if header_type == 'default' else section.first_page_header
                if not header:
                    continue

                header.is_linked_to_previous = False
                # 清除多余段落
                for i, p in enumerate(header.paragraphs):
                    if i == 0:
                        p.text = ""
                    else:
                        p._element.getparent().remove(p._element)

                p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                p.text = header_text
                p.alignment = ALIGN_MAP.get(header_align, WD_ALIGN_PARAGRAPH.CENTER)
                for run in p.runs:
                    _set_font(run, header_font, header_size)

        # ---- 页脚页码 ----
        footer_settings = settings.get("footer", {})
        pf_format = footer_settings.get("page_num_format", "第1页 / 共N页")
        footer_align = footer_settings.get("align", "居中")
        footer_font = footer_settings.get("font", "宋体")
        footer_size = cn_to_pt(footer_settings.get("size", "小五"))

        footer = section.footer
        footer.is_linked_to_previous = False

        # 清除多余段落
        for p in footer.paragraphs[1:]:
            p._element.getparent().remove(p._element)

        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = ALIGN_MAP.get(footer_align, WD_ALIGN_PARAGRAPH.CENTER)

        # 根据页码格式插入域代码
        if pf_format == "第1页 / 共N页":
            run1 = p.add_run("第 ")
            _set_font(run1, footer_font, footer_size)
            _add_page_field(p, footer_font, footer_size)
            run2 = p.add_run(" 页 / 共 ")
            _set_font(run2, footer_font, footer_size)
            _add_numpages_field(p, footer_font, footer_size)
            run3 = p.add_run(" 页")
            _set_font(run3, footer_font, footer_size)

        elif pf_format == "-1-, -2-...":
            _add_page_field(p, footer_font, footer_size)

        elif pf_format == "第1页":
            run1 = p.add_run("第 ")
            _set_font(run1, footer_font, footer_size)
            _add_page_field(p, footer_font, footer_size)
            run2 = p.add_run(" 页")
            _set_font(run2, footer_font, footer_size)

        else:  # "1, 2, 3..."
            _add_page_field(p, footer_font, footer_size)

        # 页码起始
        section.start_type = 1  # new page
        pgNumType = section._sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")} w:start="1"/>')
            section._sectPr.append(pgNumType)
        else:
            pgNumType.set(qn('w:start'), str(footer_settings.get("start_from", 1)))


def _apply_cover(doc, settings, report_title=""):
    """应用封面设置"""
    cover = settings.get("cover", {})
    if not cover.get("auto_generate"):
        return

    has_cover = False
    if doc.paragraphs:
        first_para = doc.paragraphs[0]
        if len(first_para.text) < 50 or any(kw in first_para.text for kw in ["报告", "审计", "说明", "方案"]):
            has_cover = True

    today = datetime.now()

    if has_cover:
        title_found = False
        for para in doc.paragraphs[:5]:
            if len(para.text) > 3:
                if not title_found:
                    sz = cn_to_pt(cover.get("title_size", "一号"))
                    para.alignment = ALIGN_MAP.get(cover.get("title_align", "居中"), WD_ALIGN_PARAGRAPH.CENTER)
                    for run in para.runs:
                        _set_font(run, cover.get("title_font", "宋体"), sz, cover.get("title_bold", True))
                    title_found = True
    else:
        title = report_title if report_title else "报告标题"
        title_sz = cn_to_pt(cover.get("title_size", "一号"))

        if doc.paragraphs:
            first = doc.paragraphs[0]
            first.alignment = ALIGN_MAP.get(cover.get("title_align", "居中"), WD_ALIGN_PARAGRAPH.CENTER)
            first.clear()
            run = first.add_run(title)
            _set_font(run, cover.get("title_font", "宋体"), title_sz, cover.get("title_bold", True))

            # 添加单位行
            company_para = doc.add_paragraph()
            company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            company_sz = cn_to_pt(cover.get("company_size", "三号"))
            company_run = company_para.add_run("单位名称")
            _set_font(company_run, cover.get("company_font", "宋体"), company_sz, False)

            # 添加日期行（动态日期）
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_sz = cn_to_pt(cover.get("date_size", "四号"))
            date_run = date_para.add_run(f"{today.year}年  月  日")
            _set_font(date_run, cover.get("date_font", "宋体"), date_sz, False)


def _apply_signature(doc, settings):
    """调整落款格式"""
    sig = settings.get("signature", {})
    if not sig.get("align_right"):
        return

    for para in doc.paragraphs:
        text = para.text.strip()
        if any(kw in text for kw in ["公章", "盖章", "单位", "公司", "事务所"]):
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif any(kw in text for kw in ["年", "月", "日"]) and len(text) < 20:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _apply_toc(doc, settings):
    """插入目录（插入 TOC 域）"""
    if not settings.get("toc", {}).get("include"):
        return

    # 插入"目录"标题
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("目  录")
    _set_font(run, "黑体", 16, True)

    # 插入 TOC 域
    toc_para = doc.add_paragraph()
    fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run1 = toc_para.add_run()
    run1._element.append(fld_begin)

    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2 = toc_para.add_run()
    run2._element.append(instrText)

    fld_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3 = toc_para.add_run()
    run3._element.append(fld_separate)

    run4 = toc_para.add_run("（请在 Word 中按 Ctrl+A 后按 F9 更新目录）")
    _set_font(run4, "仿宋", 9, False)

    fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5 = toc_para.add_run()
    run5._element.append(fld_end)

    # 把目录段落移动到文档最前面
    body = doc.element.body
    toc_elements = [toc_title._element, toc_para._element]
    for elem in toc_elements:
        body.remove(elem)
    first_elem = body[0] if len(body) > 0 else None
    for elem in reversed(toc_elements):
        if first_elem is not None:
            body.insert(list(body).index(first_elem), elem)
        else:
            body.append(elem)


# ========== 模板完整应用 ==========

def apply_template_to_docx(target_path, template, report_title=""):
    """将模板应用到目标 docx 文件

    Args:
        target_path: 目标 docx 文件路径
        template: 模板字典
        report_title: 报告标题（用于封面）

    Returns:
        tuple: (是否成功, 提示信息)
    """
    try:
        doc = Document(target_path)

        # 1. 应用标题样式
        for hid in [1, 2, 3]:
            key = f"heading{hid}"
            if key in template.get("headings", {}):
                _apply_heading_style(doc, key, template["headings"][key])

        # 2. 应用正文样式
        if "body" in template:
            _apply_body_style(doc, template["body"])

        # 3. 应用表格格式
        if "table" in template:
            _apply_table_format(doc, template["table"])

        # 4. 应用图片格式
        if "image" in template:
            _apply_image_format(doc, template["image"])

        # 5. 应用页眉页脚页码
        _apply_header_footer(doc, template)

        # 6. 应用封面
        _apply_cover(doc, template, report_title)

        # 7. 应用落款
        if "signature" in template:
            _apply_signature(doc, template["signature"])

        # 8. 目录
        if "toc" in template:
            _apply_toc(doc, template)

        doc.save(target_path)
        logger.info(f"模板应用成功：{target_path}")
        return True, "格式应用成功！请在 Word 中打开后按 Ctrl+A → F9 更新目录和页码。"

    except PermissionError:
        return False, "应用模板失败：文件被占用，请先关闭 Word 再试。"
    except Exception as e:
        logger.error(f"应用模板失败：{e}")
        return False, f"应用模板失败：{str(e)}"


# ========== 模板文件管理 ==========

TEMPLATES_DIR = None  # 在 GUI 初始化时设置


def init_templates_dir(base_dir):
    """初始化模板目录"""
    global TEMPLATES_DIR
    templates_dir = os.path.join(base_dir, "templates")
    if not os.path.exists(templates_dir):
        os.makedirs(templates_dir)

    for category in TEMPLATE_CATEGORIES:
        cat_dir = os.path.join(templates_dir, category)
        if not os.path.exists(cat_dir):
            os.makedirs(cat_dir)

    TEMPLATES_DIR = templates_dir
    return templates_dir


def save_template(template, name, category="通用模板"):
    """保存模板到文件"""
    if TEMPLATES_DIR is None:
        return False, "模板目录未初始化"

    cat_dir = os.path.join(TEMPLATES_DIR, category)
    if not os.path.exists(cat_dir):
        os.makedirs(cat_dir)

    template["name"] = name
    filepath = os.path.join(cat_dir, f"{name}.ftpl")

    try:
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(template, f, ensure_ascii=False, indent=2)
        logger.info(f"模板已保存：{filepath}")
        return True, f"模板已保存：{name}"
    except Exception as e:
        logger.error(f"保存模板失败：{e}")
        return False, f"保存模板失败：{str(e)}"


def load_template(filepath):
    """从文件加载模板"""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            template = json.load(f)
        return template
    except Exception as e:
        logger.error(f"加载模板失败：{e}")
        return None


def list_templates():
    """列出所有模板"""
    if TEMPLATES_DIR is None:
        return []

    result = []
    for category in TEMPLATE_CATEGORIES:
        cat_dir = os.path.join(TEMPLATES_DIR, category)
        if os.path.exists(cat_dir):
            for fname in os.listdir(cat_dir):
                if fname.endswith(".ftpl"):
                    filepath = os.path.join(cat_dir, fname)
                    try:
                        with open(filepath, "r", encoding="utf-8") as f:
                            data = json.load(f)
                        result.append({
                            "name": data.get("name", fname[:-5]),
                            "category": category,
                            "source": data.get("source", ""),
                            "filepath": filepath,
                        })
                    except Exception:
                        result.append({
                            "name": fname[:-5],
                            "category": category,
                            "source": "",
                            "filepath": filepath,
                        })

    return result


def delete_template(filepath):
    """删除模板"""
    try:
        os.remove(filepath)
        logger.info(f"模板已删除：{filepath}")
        return True, "模板已删除"
    except Exception as e:
        logger.error(f"删除模板失败：{e}")
        return False, f"删除失败：{str(e)}"


def backup_file(filepath):
    """备份文件"""
    if not os.path.exists(filepath):
        return False, "文件不存在"

    dirname = os.path.dirname(filepath)
    basename = os.path.basename(filepath)
    name, ext = os.path.splitext(basename)
    backup_name = f"{name}_备份_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
    backup_path = os.path.join(dirname, backup_name)

    try:
        shutil.copy2(filepath, backup_path)
        logger.info(f"文件已备份：{backup_path}")
        return True, backup_path
    except Exception as e:
        logger.error(f"备份失败：{e}")
        return False, f"备份失败：{str(e)}"


def detect_heading_paragraphs(doc_path):
    """检测文档中的段落结构，返回各级标题和正文数量"""
    try:
        doc = Document(doc_path)
        stats = {"h1": 0, "h2": 0, "h3": 0, "body": 0, "tables": len(doc.tables)}

        for para in doc.paragraphs:
            if para.style:
                sname = para.style.name
                if sname == "Heading 1":
                    stats["h1"] += 1
                elif sname == "Heading 2":
                    stats["h2"] += 1
                elif sname == "Heading 3":
                    stats["h3"] += 1
                elif para.text.strip():
                    stats["body"] += 1

        return stats
    except Exception as e:
        logger.error(f"检测文档结构失败：{e}")
        return None
